import os
from docx import Document
from docx.shared import Pt
import re

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    
    for line in lines:
        line = line.rstrip('\n')
        
        # Code block toggle
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue

        # Headings
        match = re.match(r'^(#{1,6})\s+(.*)', line)
        if match:
            level = len(match.group(1))
            title = match.group(2)
            doc.add_heading(title, level=level)
            continue
            
        # Horizontal rules
        if re.match(r'^---+$|^\*\*\*+$', line):
            doc.add_paragraph('---')
            continue
            
        # Lists (basic)
        list_match = re.match(r'^[\-\*]\s+(.*)', line)
        if list_match:
            p = doc.add_paragraph(list_match.group(1), style='List Bullet')
            continue
            
        # Numbered lists (basic)
        num_list_match = re.match(r'^\d+\.\s+(.*)', line)
        if num_list_match:
            p = doc.add_paragraph(num_list_match.group(1), style='List Number')
            continue
            
        # Tables (not handled perfectly, but let's try to just put text for now)
        if line.startswith('|'):
            # Basic table row
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            continue

        # Regular paragraph
        if line.strip():
            # Handle basic bold/italic within line
            p = doc.add_paragraph()
            # Simple bold/italic parsing (not exhaustive)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
        else:
            # Empty line
            pass

    doc.save(docx_path)

files = [
    "part_1_introduction.md",
    "part_2_literature_review.md",
    "part_3_methodology.md",
    "part_4_results.md",
    "part_5_discussion_conclusion.md"
]

for f in files:
    out = f.replace(".md", ".docx")
    print(f"Converting {f} to {out}...")
    convert_md_to_docx(f, out)
    print("Done.")
